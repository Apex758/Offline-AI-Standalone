import io
import csv
import json
import base64
from typing import Any, Dict, List, Union
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def add_table_borders(table):
    """Add borders to a table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def parse_lesson_content_for_export(text: str) -> List[Dict[str, Any]]:
    """Parse lesson plan text into structured sections for export"""
    if not text:
        return []
    
    lines = text.split('\n')
    sections = []
    current_section = None
    details = []
    
    for line in lines:
        trimmed = line.strip()
        
        if not trimmed:
            continue
        
        # Collect basic details for grid
        match = re.match(r'^\*\*(Grade Level|Subject|Strand|Topic|Duration|Date):\*\* (.+)', trimmed)
        if match:
            label, value = match.groups()
            details.append({'label': label, 'value': value})
            if len(details) == 6:
                sections.append({'type': 'details_grid', 'items': details.copy()})
                details = []
            continue
        
        # Skip main lesson title
        if re.match(r'^\*\*Lesson Plan:', trimmed):
            continue
        
        # Section headings (surrounded by **)
        match = re.match(r'^\*\*(.+)\*\*$', trimmed)
        if match:
            title = match.group(1)
            current_section = {'type': 'section', 'title': title, 'content': []}
            sections.append(current_section)
            continue
        
        # Field labels (start with ** but don't end with **)
        if re.match(r'^\*\*[^*]+:\*\*', trimmed) or re.match(r'^\*\*[^*]+:$', trimmed):
            title = re.sub(r'^\*\*', '', trimmed).replace('**', '').replace(':', '').strip()
            current_section = {'type': 'field', 'title': title, 'content': []}
            sections.append(current_section)
            continue
        
        # Bullet points with + (nested)
        match = re.match(r'^\s*\+\s+(.+)', trimmed)
        if match:
            content = match.group(1)
            if current_section:
                current_section['content'].append({'type': 'nested_bullet', 'text': content})
            continue
        
        # Regular bullet points
        match = re.match(r'^\s*\*\s+(.+)', trimmed)
        if match:
            if not trimmed.startswith('**'):
                content = match.group(1)
                if current_section:
                    current_section['content'].append({'type': 'bullet', 'text': content})
                continue
        
        # Numbered items
        match = re.match(r'^(\d+)\.\s*(.+)', trimmed)
        if match:
            number, content = match.groups()
            if current_section:
                current_section['content'].append({'type': 'numbered', 'number': number, 'text': content})
            continue
        
        # Regular paragraphs
        if current_section:
            current_section['content'].append({'type': 'paragraph', 'text': trimmed})
    
    return sections


def add_qr_to_docx(doc, qr_bytes: bytes, title_text: str = ''):
    """Add QR code image to the header area of the DOCX.

    Creates a two-column header: left = title area, right = QR code.
    """
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Create a table in header: left cell = title area, right cell = QR
    table = header.add_table(rows=1, cols=2, width=Cm(19))
    table.columns[0].width = Cm(14)
    table.columns[1].width = Cm(5)

    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Left cell: optional title text
    if title_text:
        left_cell = table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        run = left_para.add_run(title_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(156, 163, 175)

    # Right cell: QR code
    qr_cell = table.cell(0, 1)
    qr_paragraph = qr_cell.paragraphs[0]
    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = qr_paragraph.add_run()
    run.add_picture(io.BytesIO(qr_bytes), width=Cm(2), height=Cm(2))


def add_alignment_markers_to_html(html: str) -> str:
    """Inject alignment marker CSS + elements into HTML for PDF export.

    Adds 3 filled squares at top-left, bottom-left, bottom-right corners.
    """
    marker_css = """
    .alignment-marker {
        position: fixed;
        width: 5mm;
        height: 5mm;
        background-color: black;
        z-index: 9999;
    }
    .marker-top-left { top: 5mm; left: 5mm; }
    .marker-bottom-left { bottom: 5mm; left: 5mm; }
    .marker-bottom-right { bottom: 5mm; right: 5mm; }
    """
    marker_elements = """
    <div class="alignment-marker marker-top-left"></div>
    <div class="alignment-marker marker-bottom-left"></div>
    <div class="alignment-marker marker-bottom-right"></div>
    """

    # Inject CSS before </head> or at start of <style>
    if '</head>' in html:
        html = html.replace('</head>',
                            f'<style>{marker_css}</style></head>')
    elif '<style>' in html:
        html = html.replace('<style>', f'<style>{marker_css}')
    else:
        html = f'<style>{marker_css}</style>' + html

    # Inject marker elements after <body> or at start
    if '<body' in html:
        # Find end of body tag
        body_end = html.index('>', html.index('<body')) + 1
        html = html[:body_end] + marker_elements + html[body_end:]
    else:
        html = marker_elements + html

    return html


def inject_qr_into_html(html: str, qr_base64: str) -> str:
    """Inject QR code image into the top-right corner of HTML for PDF export.

    Args:
        html: The document HTML
        qr_base64: Base64-encoded PNG of the QR code
    """
    qr_css = """
    .scan-qr-code {
        position: fixed;
        top: 5mm;
        right: 5mm;
        width: 20mm;
        height: 20mm;
        z-index: 9999;
    }
    """
    qr_element = f'<img class="scan-qr-code" src="data:image/png;base64,{qr_base64}" />'

    if '</head>' in html:
        html = html.replace('</head>',
                            f'<style>{qr_css}</style></head>')
    elif '<style>' in html:
        html = html.replace('<style>', f'<style>{qr_css}')

    if '<body' in html:
        body_end = html.index('>', html.index('<body')) + 1
        html = html[:body_end] + qr_element + html[body_end:]
    else:
        html = qr_element + html

    return html


def export_to_docx(data: Union[str, Dict, List], title: str = "Export") -> bytes:
    """
    Export to DOCX format.
    Uses rawHtml if provided for perfect visual consistency with PDF and screen.
    """
    try:
        # Get data
        if isinstance(data, dict):
            raw_html = data.get('rawHtml', '')
            content = data.get('content', '')
            form_data = data.get('formData', {})
            accent_color = data.get('accentColor', '#3B82F6')
        else:
            raw_html = ''
            content = str(data)
            form_data = {}
            accent_color = '#3B82F6'
        
        # ✅ PRIORITY: Use HTML parser for perfect consistency
        if raw_html:
            return export_to_docx_from_html(raw_html, accent_color, form_data)
        
        # Fallback to simple text-based DOCX
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(content)
        
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
         
    except Exception as e:
        raise RuntimeError(f"Failed to export DOCX: {e}")



def export_to_pdf(data: Union[str, Dict, List], title: str = "Export") -> bytes:
    """
    Export to PDF using HTML.
    If rawHtml is provided, uses it directly for perfect consistency.
    """
    import os
    import sys
    if sys.platform == "win32":
        script_dir = os.path.dirname(os.path.abspath(__file__))
        gtk_bin = os.path.join(script_dir, "bin")
        if not os.path.isdir(gtk_bin):
            parent_dir = os.path.dirname(script_dir)
            backend_bundle_bin = os.path.join(parent_dir, "backend-bundle", "bin")
            if os.path.isdir(backend_bundle_bin):
                gtk_bin = backend_bundle_bin
        
        if os.path.isdir(gtk_bin):
            os.environ["WEASYPRINT_DLL_DIRECTORIES"] = gtk_bin
            if hasattr(os, "add_dll_directory"):
                try:
                    os.add_dll_directory(gtk_bin)
                except:
                    pass
            os.environ["PATH"] = gtk_bin + os.pathsep + os.environ.get("PATH", "")
    
    from weasyprint import HTML
    try:
        # ✅ PRIORITY: Use rawHtml if provided (perfect consistency)
        if isinstance(data, dict) and data.get("rawHtml"):
            html = data["rawHtml"]
            pdf_bytes = HTML(string=html).write_pdf()
            return pdf_bytes

        # Fallback to legacy HTML generation
        if isinstance(data, dict):
            content = data.get('content', '')
            form_data = data.get('formData', {})
            accent_color = data.get('accentColor', '#3B82F6')
        else:
            content = str(data)
            form_data = {}
            accent_color = '#3B82F6'

        # Build basic HTML (legacy fallback)
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: 'Segoe UI', sans-serif;
                    line-height: 1.6;
                    color: #374151;
                    margin: 2cm;
                }}
                pre {{
                    white-space: pre-wrap;
                    font-family: inherit;
                }}
            </style>
        </head>
        <body>
            <h1 style="color: {accent_color};">{title}</h1>
            <pre>{content}</pre>
        </body>
        </html>
        """

        pdf_bytes = HTML(string=html).write_pdf()
        return pdf_bytes

    except Exception as e:
        raise RuntimeError(f"Failed to export PDF: {e}")


async def export_to_pdf_async(data, title: str = "Export") -> bytes:
    """Async wrapper for export_to_pdf — runs in a thread to avoid blocking the event loop."""
    import asyncio
    return await asyncio.get_running_loop().run_in_executor(None, lambda: export_to_pdf(data, title))


# ---------------------------------------------------------------------------
# Lesson-plan DOCX builder (Phase 2 of WYSIWYG lock-in)
#
# Mirrors frontend/src/utils/lessonExportSpec.ts. Keep these constants in sync
# if the frontend spec changes. The page margin and table column widths are
# also embedded in the generated HTML (@page rule + <colgroup>), so the PDF
# pipeline (WeasyPrint) and the DOCX pipeline produce matching output.
# ---------------------------------------------------------------------------
LESSON_PAGE_MARGIN_MM = 10.0
LESSON_BODY_PT = 11
LESSON_SMALL_PT = 9
LESSON_MICRO_PT = 8
LESSON_H1_PT = 18


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply solid background shading to a python-docx table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#').upper())
    tcPr.append(shd)


def _extract_bg_hex(style: str):
    """Pull a background colour out of an inline style attribute.

    Accepts #RRGGBB and #RRGGBBAA (alpha is dropped because Word shading is
    opaque). Returns None if no usable colour is found.
    """
    if not style:
        return None
    m = re.search(r'background\s*:\s*([^;]+)', style)
    if not m:
        return None
    val = m.group(1).strip()
    h = re.match(r'#([0-9a-fA-F]{6})(?:[0-9a-fA-F]{2})?', val)
    if h:
        return h.group(1)
    return None


def _parse_colgroup_widths(table_tag) -> list:
    """Read explicit percentage widths from a <colgroup>/<col> block."""
    cg = table_tag.find('colgroup')
    if not cg:
        return []
    widths = []
    for col in cg.find_all('col'):
        s = col.get('style', '') or ''
        m = re.search(r'width\s*:\s*([\d.]+)%', s)
        if m:
            widths.append(float(m.group(1)))
    return widths


def _render_cell_content(cell, source_tag, default_pt: int,
                         body_color: tuple, accent: tuple) -> None:
    """Walk a BeautifulSoup cell and emit paragraphs into a docx cell.

    Strips the empty paragraph that python-docx auto-creates so the cell
    starts clean. Recognises:
      - <ul>/<li>            -> bulleted lines (ASCII '-' bullet)
      - small bold <div>     -> uppercase label rendered in accent colour
      - italic <div>         -> italic line
      - plain text/divs      -> body paragraphs
    """
    # Drop the default empty paragraph
    if cell.paragraphs and not cell.paragraphs[0].runs:
        p0 = cell.paragraphs[0]._element
        p0.getparent().remove(p0)

    def add_paragraph(text, bold=False, italic=False, size_pt=default_pt,
                      color=body_color, bullet=False):
        text = (text or '').strip()
        if not text:
            return
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if bullet:
            lead = p.add_run('-  ')
            lead.font.size = Pt(size_pt)
            lead.font.color.rgb = RGBColor(*color)
        run = p.add_run(text)
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*color)

    def style_flags(style: str):
        small = 'font-size:10px' in style or 'font-size:11px' in style
        is_label = ('font-weight:700' in style and 'text-transform:uppercase' in style)
        is_bold = 'font-weight:700' in style or 'font-weight:600' in style
        is_italic = 'font-style:italic' in style
        return small, is_label, is_bold, is_italic

    def walk(node):
        for child in node.children:
            name = getattr(child, 'name', None)
            # NavigableString (raw text)
            if name is None:
                text = str(child).strip()
                if text:
                    add_paragraph(text)
                continue
            if name == 'br':
                continue
            if name == 'ul':
                for li in child.find_all('li', recursive=False):
                    add_paragraph(li.get_text(' ', strip=True), bullet=True)
                continue
            if name == 'span':
                text = child.get_text(' ', strip=True)
                if text:
                    add_paragraph(text)
                continue
            if name == 'div':
                style = child.get('style', '') or ''
                small, is_label, is_bold, is_italic = style_flags(style)
                has_block_kids = any(
                    getattr(c, 'name', None) in ('ul', 'div')
                    for c in child.children
                )
                if has_block_kids:
                    # Emit any direct (non-block) text first, then recurse.
                    direct_text = ' '.join(
                        str(c).strip() for c in child.children
                        if getattr(c, 'name', None) not in ('ul', 'div') and str(c).strip()
                    )
                    if direct_text:
                        snippet = BeautifulSoup(direct_text, 'html.parser').get_text(' ', strip=True)
                        if snippet:
                            add_paragraph(
                                snippet,
                                bold=is_bold,
                                italic=is_italic,
                                size_pt=LESSON_SMALL_PT if small else default_pt,
                                color=accent if is_label else body_color,
                            )
                    walk(child)
                else:
                    text = child.get_text(' ', strip=True)
                    add_paragraph(
                        text,
                        bold=is_bold,
                        italic=is_italic,
                        size_pt=LESSON_SMALL_PT if small else default_pt,
                        color=accent if is_label else body_color,
                    )
                continue
            # Fallback: any other tag, take its text
            text = child.get_text(' ', strip=True) if hasattr(child, 'get_text') else str(child)
            if text:
                add_paragraph(text)

    walk(source_tag)

    # Guarantee the cell has at least one paragraph (Word requires it)
    if not cell.paragraphs:
        cell.add_paragraph()


def _build_one_lesson_table(table_tag, doc, accent, body_color, usable_in: float) -> None:
    """Convert one <table> in the lesson HTML into a real Word table.

    Honours <colgroup> percentage widths and rowspan/colspan via cell merging.
    """
    pct_widths = _parse_colgroup_widths(table_tag)
    if not pct_widths:
        first_row = table_tag.find('tr')
        ncols = (
            sum(int(c.get('colspan', 1)) for c in first_row.find_all(['td', 'th']))
            if first_row else 1
        )
        pct_widths = [100.0 / ncols] * ncols
    ncols = len(pct_widths)

    rows = []
    thead = table_tag.find('thead')
    if thead:
        rows.extend(thead.find_all('tr'))
    tbody = table_tag.find('tbody')
    if tbody:
        rows.extend(tbody.find_all('tr'))
    if not rows:
        rows = table_tag.find_all('tr')
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.autofit = False
    table.allow_autofit = False
    add_table_borders(table)

    # Lock per-column widths from the colgroup
    for i, pct in enumerate(pct_widths):
        col_in = usable_in * (pct / 100.0)
        table.columns[i].width = Inches(col_in)
        for row in table.rows:
            row.cells[i].width = Inches(col_in)

    for r_idx, tr in enumerate(rows):
        col_idx = 0
        for src_cell in tr.find_all(['td', 'th'], recursive=False):
            if col_idx >= ncols:
                break
            colspan = int(src_cell.get('colspan', 1))
            target = table.cell(r_idx, col_idx)
            if colspan > 1:
                end = min(col_idx + colspan - 1, ncols - 1)
                target = target.merge(table.cell(r_idx, end))

            style = src_cell.get('style', '') or ''
            bg = _extract_bg_hex(style)
            if bg:
                _set_cell_shading(target, bg)

            _render_cell_content(target, src_cell, LESSON_BODY_PT, body_color, accent)

            # Header cells / explicitly bold cells: force bold on every run
            is_header = (
                src_cell.name == 'th'
                or 'font-weight:700' in style
                or 'font-weight:600' in style
            )
            if is_header:
                for p in target.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

            col_idx += colspan


def _build_lesson_plan_docx(root_div, doc, rgb: tuple) -> None:
    """Top-level builder for the lesson-plan DOCX export.

    Sets narrow page margins (matching the PDF @page rule), renders the title,
    then walks each <table> child of the lesson root and rebuilds it as a real
    Word table with locked column widths.
    """
    section = doc.sections[0]
    margin_cm = LESSON_PAGE_MARGIN_MM / 10.0
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)

    body_color = (34, 34, 34)
    accent = rgb

    # A4 width (21cm) minus left+right margins, in inches
    usable_cm = 21.0 - 2 * margin_cm
    usable_in = usable_cm / 2.54

    h1 = root_div.find('h1')
    if h1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(h1.get_text(strip=True))
        run.font.size = Pt(LESSON_H1_PT)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*accent)

    # Walk children in document order so tables and post-table notes interleave
    for child in root_div.find_all(recursive=False):
        if child.name == 'table':
            _build_one_lesson_table(child, doc, accent, body_color, usable_in)
            doc.add_paragraph()  # spacer between tables
        elif child.name == 'div':
            # Curriculum references / footer-style blocks: render as flat text
            text = child.get_text(' ', strip=True)
            if text:
                p = doc.add_paragraph(text)
                for r in p.runs:
                    r.font.size = Pt(LESSON_SMALL_PT)
                    r.font.color.rgb = RGBColor(*body_color)
        elif child.name == 'h1':
            continue  # already handled


def _attach_export_footer(doc, form_data: dict, rgb: tuple) -> None:
    """Shared footer (generated-on text + OECS logo) for all DOCX exports."""
    import os
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    text_run = footer_p.add_run(f"Generated on {form_data.get('date', 'N/A')}")
    text_run.font.size = Pt(9)
    text_run.font.color.rgb = RGBColor(156, 163, 175)

    logo_paths = [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'build', 'OECS.png'),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'frontend', 'public', 'OECS.png'),
    ]
    for logo_path in logo_paths:
        if os.path.isfile(logo_path):
            footer_p.add_run("    ")
            logo_run = footer_p.add_run()
            logo_run.add_picture(logo_path, height=Pt(30))
            break


def export_to_docx_from_html(html: str, accent_color: str, form_data: dict) -> bytes:
    """
    Convert HTML to DOCX while preserving formatting.
    This ensures DOCX matches PDF and screen display exactly.
    """
    doc = Document()
    soup = BeautifulSoup(html, 'html.parser')
    rgb = hex_to_rgb(accent_color)

    # NEW: lesson-plan exports use a dedicated table-preserving builder so the
    # DOCX honours the same locked column widths and font sizes as the PDF.
    lesson_root = soup.find('div', id='lesson-plan-html-export')
    if lesson_root is not None:
        _build_lesson_plan_docx(lesson_root, doc, rgb)
        _attach_export_footer(doc, form_data, rgb)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    
    # Parse header section (gradient banner)
    # Find the header div by looking for the subject badge
    header_section = soup.find('div', style=lambda x: x and 'gradient' in str(x) if x else False)
    
    if not header_section:
        # Try finding by structure
        for div in soup.find_all('div'):
            if div.find('span') and div.find('h1'):
                header_section = div
                break
    
    if header_section:
        # Add subject badge
        badge = header_section.find('span')
        if badge:
            p = doc.add_paragraph()
            run = p.add_run(badge.text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            # Simulate badge background with shading
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Spacing
        
        # Add main title
        title_h1 = header_section.find('h1')
        if title_h1:
            p = doc.add_paragraph()
            run = p.add_run(title_h1.text)
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(*rgb)
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Spacing
    
    # Parse main content
    # Get the content div (after header)
    content_divs = soup.find_all('div', recursive=False)
    
    # Find content div (usually the one after header with margin-top)
    content_div = None
    for div in content_divs:
        style = div.get('style', '')
        if 'margin-top: 2rem' in style or div.find(['h2', 'h3', 'p']):
            content_div = div
            break
    
    if not content_div:
        # Fallback: get body content
        content_div = soup.find('body')
    
    if content_div:
        # Process images embedded as base64 data URIs
        for img in content_div.find_all('img'):
            src = img.get('src', '')
            if src.startswith('data:image/'):
                try:
                    # Extract base64 data from data URI
                    header, b64_data = src.split(',', 1)
                    image_bytes = base64.b64decode(b64_data)
                    image_stream = io.BytesIO(image_bytes)
                    # Determine width from style if available
                    style = img.get('style', '')
                    width = Inches(3)  # default
                    width_match = re.search(r'(?:width|max-width)\s*:\s*(\d+)rem', style)
                    if width_match:
                        rem_val = int(width_match.group(1))
                        width = Inches(min(rem_val * 0.5, 5.5))  # cap at page width
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(image_stream, width=width)
                    doc.add_paragraph()  # spacing after image
                except Exception as img_err:
                    # If image fails, skip silently rather than breaking the export
                    pass

        # Process each element
        for element in content_div.find_all(['h2', 'h3', 'div', 'p'], recursive=True):
            style = element.get('style', '')
            text = element.get_text(strip=True)

            if not text:
                continue

            # Section headings (h2)
            if element.name == 'h2':
                heading = doc.add_heading(text, level=1)
                heading.runs[0].font.color.rgb = RGBColor(*rgb)
                heading.runs[0].font.size = Pt(18)
                
            # Question headings (h3)
            elif element.name == 'h3':
                heading = doc.add_heading(text, level=2)
                heading.runs[0].font.color.rgb = RGBColor(
                    int(rgb[0] * 0.8),
                    int(rgb[1] * 0.8),
                    int(rgb[2] * 0.8)
                )
                heading.runs[0].font.size = Pt(14)
                
            # Answer options (divs with A), B), etc.)
            elif element.name == 'div' and re.match(r'^[A-D]\)', text):
                p = doc.add_paragraph()
                letter_run = p.add_run(text[:2] + ' ')
                letter_run.font.bold = True
                letter_run.font.color.rgb = RGBColor(*rgb)
                content_run = p.add_run(text[2:].strip())
                p.paragraph_format.left_indent = Inches(0.25)
                
            # Bullet points (divs with •)
            elif element.name == 'div' and '•' in text[:2]:
                content = text.replace('•', '').strip()
                p = doc.add_paragraph(content, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                
            # Numbered lists (divs with numbers)
            elif element.name == 'div' and re.match(r'^\d+\.', text):
                content = re.sub(r'^\d+\.\s*', '', text)
                p = doc.add_paragraph(content, style='List Number')
                p.paragraph_format.left_indent = Inches(0.25)
                
            # Regular paragraphs
            elif element.name == 'p':
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)
    
    # Add footer with OECS logo
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add "Generated on" text as left-aligned tab, then logo right-aligned
    text_run = footer_p.add_run(f"Generated on {form_data.get('date', 'N/A')}")
    text_run.font.size = Pt(9)
    text_run.font.color.rgb = RGBColor(156, 163, 175)

    # Add OECS logo to footer
    import os
    logo_paths = [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'build', 'OECS.png'),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'frontend', 'public', 'OECS.png'),
    ]
    for logo_path in logo_paths:
        if os.path.isfile(logo_path):
            footer_p.add_run("    ")
            logo_run = footer_p.add_run()
            logo_run.add_picture(logo_path, height=Pt(30))
            break
    
    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def format_section_content_html(content: List[Dict]) -> str:
    """Format section content as HTML"""
    html = ''
    
    # Group consecutive bullets and numbered items
    current_list_type = None
    
    for item in content:
        item_type = item['type']
        
        if item_type == 'bullet' or item_type == 'nested_bullet':
            if current_list_type != 'bullet':
                if current_list_type:
                    html += '</ul>' if current_list_type == 'bullet' else '</ol>'
                html += '<ul class="bullet-list">'
                current_list_type = 'bullet'
            
            css_class = 'nested-bullet' if item_type == 'nested_bullet' else ''
            html += f'<li class="{css_class}">{item["text"]}</li>'
        
        elif item_type == 'numbered':
            if current_list_type != 'numbered':
                if current_list_type:
                    html += '</ul>' if current_list_type == 'bullet' else '</ol>'
                html += '<ol class="numbered-list">'
                current_list_type = 'numbered'
            
            html += f'<li>{item["text"]}</li>'
        
        elif item_type == 'paragraph':
            if current_list_type:
                html += '</ul>' if current_list_type == 'bullet' else '</ol>'
                current_list_type = None
            
            html += f'<p>{item["text"]}</p>'
    
    # Close any open lists
    if current_list_type:
        html += '</ul>' if current_list_type == 'bullet' else '</ol>'
    
    return html


def export_to_csv(data: Union[List[Dict], List[List], Dict, List], title: str = "Export") -> bytes:
    """
    Export data to CSV format.
    """
    try:
        buf = io.StringIO()
        writer = csv.writer(buf)
        if isinstance(data, list):
            if all(isinstance(row, dict) for row in data):
                # Write header
                if data:
                    writer.writerow(data[0].keys())
                    for row in data:
                        writer.writerow(row.values())
            elif all(isinstance(row, list) for row in data):
                for row in data:
                    writer.writerow(row)
            else:
                for item in data:
                    writer.writerow([item])
        elif isinstance(data, dict):
            writer.writerow(data.keys())
            writer.writerow(data.values())
        else:
            writer.writerow([str(data)])
        return buf.getvalue().encode("utf-8")
    except Exception as e:
        raise RuntimeError(f"Failed to export CSV: {e}")


def export_to_json(data: Any, title: str = "Export") -> bytes:
    """
    Export data to JSON format.
    """
    try:
        return json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8")
    except Exception as e:
        raise RuntimeError(f"Failed to export JSON: {e}")


def export_to_markdown(data: Union[str, Dict, List], title: str = "Export") -> bytes:
    """
    Export data to Markdown format.
    """
    try:
        md = f"# {title}\n\n"
        if isinstance(data, str):
            md += data
        elif isinstance(data, dict):
            for k, v in data.items():
                md += f"- **{k}**: {v}\n"
        elif isinstance(data, list):
            for item in data:
                md += f"- {item}\n"
        else:
            md += str(data)
        return md.encode("utf-8")
    except Exception as e:
        raise RuntimeError(f"Failed to export Markdown: {e}")


EXPORT_FORMATTERS = {
    "pdf": export_to_pdf,
    "docx": export_to_docx,
    "csv": export_to_csv,
    "json": export_to_json,
    "md": export_to_markdown,
    "markdown": export_to_markdown,
}